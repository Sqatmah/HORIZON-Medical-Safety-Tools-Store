from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils import timezone

BRAND_TEAL = RGBColor(0x1A, 0x4B, 0x4C)
BRAND_ACCENT = RGBColor(0x00, 0xA8, 0xCC)


def generate_correspondence_docx(corr):
    doc = Document()

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('مؤسسة الابتكار التقني')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BRAND_TEAL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run('Tech Innovation Establishment')
    run2.font.size = Pt(11)
    run2.font.color.rgb = BRAND_ACCENT

    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f"التاريخ: {timezone.now().strftime('%Y-%m-%d')}").font.size = Pt(11)

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = 'عرض سعر' if corr.purpose == 'quotation' else 'تأكيد طلبية'
    run3 = title.add_run(title_text)
    run3.font.color.rgb = BRAND_TEAL

    doc.add_paragraph(f"السادة / {corr.recipient_name}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"البريد الإلكتروني: {corr.recipient_email}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"نوع الجهة: {corr.get_recipient_type_display()}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    if corr.purpose == 'quotation':
        body = doc.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        body.add_run(
            f"يسر مؤسسة الابتكار التقني تقديم عرض السعر التالي بناءً على طلبكم:"
        )
        if corr.quantity:
            doc.add_paragraph(f"الكمية المطلوبة: {corr.quantity}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if corr.amount:
            amt = doc.add_paragraph()
            amt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            amt_run = amt.add_run(f"إجمالي قيمة العرض: {corr.amount} ريال سعودي")
            amt_run.font.bold = True
            amt_run.font.size = Pt(13)
    else:
        body = doc.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        body.add_run("نؤكد لكم استلام واعتماد الطلبية التالية:")
        if corr.quantity:
            doc.add_paragraph(f"الكمية: {corr.quantity}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if corr.amount:
            amt = doc.add_paragraph()
            amt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            amt_run = amt.add_run(f"القيمة الإجمالية: {corr.amount} ريال سعودي")
            amt_run.font.bold = True
            amt_run.font.size = Pt(13)

    if corr.notes:
        doc.add_paragraph()
        notes_title = doc.add_paragraph()
        notes_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        notes_title.add_run("ملاحظات:").font.bold = True
        doc.add_paragraph(corr.notes).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph('─' * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("مع خالص التقدير — إدارة مؤسسة الابتكار التقني").font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer